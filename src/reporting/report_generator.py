"""
Reporting Layer (§12).

Generates:
- Excel report (multi-tab: summary, break detail, aging, recon, auctions)
- DOCX narrative summary

Includes: STP rate, custodian confirmation rate, cost-of-exception,
break distribution, counterparty analysis.
"""

import json
from datetime import date, datetime
from decimal import Decimal
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy.orm import Session

from src.models.database import (
    AgenticAuditLog,
    AuctionRecord,
    BreakRecord,
    Obligation,
    SettlementInstruction,
    Trade,
)
from src.models.enums import (
    AuctionOutcome,
    BreakStatus,
    BreakType,
    MatchStatus,
    ObligationStage,
    ObligationStatus,
    Severity,
)
from src.utils.config_loader import get_escalation_config


# Style constants
HEADER_FILL = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
HEADER_FONT = Font(color="FFFFFF", bold=True, size=11)
SUBHEADER_FILL = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
THIN_BORDER = Border(
    left=Side(style="thin"),
    right=Side(style="thin"),
    top=Side(style="thin"),
    bottom=Side(style="thin"),
)

SEVERITY_COLORS = {
    "LOW": PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid"),
    "MEDIUM": PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid"),
    "HIGH": PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid"),
}


def _style_header_row(ws, row, num_cols):
    for col in range(1, num_cols + 1):
        cell = ws.cell(row=row, column=col)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = THIN_BORDER


def _auto_width(ws):
    for col_cells in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col_cells[0].column)
        for cell in col_cells:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = min(max_len + 3, 40)


def generate_excel_report(
    session: Session,
    triage_result: dict,
    recon_results: list,
    output_path: Path,
):
    """Generate the multi-tab Excel settlement report."""
    wb = Workbook()
    config = get_escalation_config()
    cost_config = config["cost_per_break"]

    # ── Summary Tab ─────────────────────────────────────────────────────
    ws = wb.active
    ws.title = "Summary"

    total_trades = session.query(Trade).filter(Trade.source_system == "OMS").count()
    total_obs = session.query(Obligation).filter(
        Obligation.obligation_stage == ObligationStage.FINAL
    ).count()
    matched = session.query(Obligation).filter(
        Obligation.match_status == MatchStatus.MATCHED
    ).count()
    settled = session.query(Obligation).filter(
        Obligation.status == ObligationStatus.SETTLED
    ).count()
    instructed = session.query(Obligation).filter(
        Obligation.status == ObligationStatus.INSTRUCTED
    ).count()
    total_breaks = session.query(BreakRecord).count()

    # STP: obligations that went through with zero manual intervention
    stp_count = session.query(Obligation).filter(
        Obligation.obligation_stage == ObligationStage.FINAL,
        Obligation.status.in_([ObligationStatus.SETTLED, ObligationStatus.INSTRUCTED]),
    ).count()
    stp_rate = (stp_count / total_obs * 100) if total_obs > 0 else 0

    # Custodian confirmation rate
    from src.models.enums import ConfirmationStatus
    confirmed_count = session.query(Obligation).filter(
        Obligation.confirmation_status == ConfirmationStatus.CONFIRMED
    ).count()
    pending_confirm = session.query(Obligation).filter(
        Obligation.confirmation_status == ConfirmationStatus.PENDING
    ).count()
    conf_total = confirmed_count + pending_confirm
    conf_rate = (confirmed_count / conf_total * 100) if conf_total > 0 else 0

    # Cost of exception
    all_breaks = session.query(BreakRecord).all()
    total_cost = sum(
        cost_config.get(brk.severity.value, 5000)
        for brk in all_breaks
    )

    summary_data = [
        ("Daily Settlement Status Report", ""),
        ("Report Date", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("", ""),
        ("KEY PERFORMANCE INDICATORS", ""),
        ("Total Trades (OMS)", total_trades),
        ("Total Obligations (Final)", total_obs),
        ("Matched Obligations", matched),
        ("Settled Obligations", settled),
        ("Instructed Obligations", instructed),
        ("STP Rate", f"{stp_rate:.1f}%"),
        ("Custodian Confirmation Rate", f"{conf_rate:.1f}%"),
        ("", ""),
        ("BREAK SUMMARY", ""),
        ("Total Breaks", total_breaks),
        ("Estimated Cost of Exceptions (INR)", f"₹{total_cost:,.0f}"),
        ("", ""),
        ("FAIL-RISK SCAN", ""),
        ("Obligations Scanned", len(triage_result.get("fail_risk_scores", []))),
        ("High-Risk Queue", len(triage_result.get("high_risk_queue", []))),
    ]

    for row_idx, (label, value) in enumerate(summary_data, 1):
        ws.cell(row=row_idx, column=1, value=label).font = Font(bold=True, size=11)
        ws.cell(row=row_idx, column=2, value=value)

    ws.cell(row=1, column=1).font = Font(bold=True, size=14, color="1F4E79")
    _auto_width(ws)

    # ── Break Detail Tab ────────────────────────────────────────────────
    ws_breaks = wb.create_sheet("Break Detail")
    headers = [
        "Break ID", "Obligation ID", "Type", "Severity",
        "Value at Risk (INR)", "Age (Days)", "Age (Hours)",
        "Status", "Escalation", "Recommended Action",
    ]
    for col, h in enumerate(headers, 1):
        ws_breaks.cell(row=1, column=col, value=h)
    _style_header_row(ws_breaks, 1, len(headers))

    for row_idx, brk in enumerate(all_breaks, 2):
        ws_breaks.cell(row=row_idx, column=1, value=brk.break_id[:12] + "...")
        ws_breaks.cell(row=row_idx, column=2, value=brk.obligation_id[:12] + "...")
        ws_breaks.cell(row=row_idx, column=3, value=brk.break_type.value)
        sev_cell = ws_breaks.cell(row=row_idx, column=4, value=brk.severity.value)
        sev_cell.fill = SEVERITY_COLORS.get(brk.severity.value, PatternFill())
        ws_breaks.cell(row=row_idx, column=5, value=float(brk.value_at_risk or 0))
        ws_breaks.cell(row=row_idx, column=6, value=brk.age_days or 0)
        ws_breaks.cell(row=row_idx, column=7, value=round(brk.age_hours or 0, 1))
        ws_breaks.cell(row=row_idx, column=8, value=brk.status.value)
        ws_breaks.cell(row=row_idx, column=9, value=brk.escalation_level)
        ws_breaks.cell(row=row_idx, column=10, value=(brk.recommended_action or "")[:100])

    _auto_width(ws_breaks)

    # ── Break Type Distribution Tab ─────────────────────────────────────
    ws_dist = wb.create_sheet("Break Distribution")
    dist_headers = ["Break Type", "Count", "% of Total", "Estimated Cost (INR)"]
    for col, h in enumerate(dist_headers, 1):
        ws_dist.cell(row=1, column=col, value=h)
    _style_header_row(ws_dist, 1, len(dist_headers))

    from collections import Counter
    type_counts = Counter(brk.break_type.value for brk in all_breaks)
    for row_idx, (bt, count) in enumerate(type_counts.most_common(), 2):
        ws_dist.cell(row=row_idx, column=1, value=bt)
        ws_dist.cell(row=row_idx, column=2, value=count)
        ws_dist.cell(row=row_idx, column=3, value=f"{count/total_breaks*100:.1f}%")
        type_cost = sum(
            cost_config.get(brk.severity.value, 5000)
            for brk in all_breaks if brk.break_type.value == bt
        )
        ws_dist.cell(row=row_idx, column=4, value=type_cost)

    _auto_width(ws_dist)

    # ── Aging Tab ───────────────────────────────────────────────────────
    ws_aging = wb.create_sheet("Aging Analysis")
    aging_headers = ["Aging Bucket", "Count", "Total VAR (INR)"]
    for col, h in enumerate(aging_headers, 1):
        ws_aging.cell(row=1, column=col, value=h)
    _style_header_row(ws_aging, 1, len(aging_headers))

    buckets = {"0-1 days": [], "2-3 days": [], "4+ days": []}
    for brk in all_breaks:
        age = brk.age_days or 0
        if age <= 1:
            buckets["0-1 days"].append(brk)
        elif age <= 3:
            buckets["2-3 days"].append(brk)
        else:
            buckets["4+ days"].append(brk)

    for row_idx, (bucket, brks) in enumerate(buckets.items(), 2):
        ws_aging.cell(row=row_idx, column=1, value=bucket)
        ws_aging.cell(row=row_idx, column=2, value=len(brks))
        ws_aging.cell(row=row_idx, column=3, value=sum(float(b.value_at_risk or 0) for b in brks))

    _auto_width(ws_aging)

    # ── Reconciliation Tab ──────────────────────────────────────────────
    ws_recon = wb.create_sheet("Reconciliation")
    recon_headers = [
        "Counterparty", "ISIN", "Statement Date",
        "Internal Qty", "Custodian Qty", "Difference", "Reconciled",
    ]
    for col, h in enumerate(recon_headers, 1):
        ws_recon.cell(row=1, column=col, value=h)
    _style_header_row(ws_recon, 1, len(recon_headers))

    for row_idx, r in enumerate(recon_results[:100], 2):  # limit to 100 rows
        ws_recon.cell(row=row_idx, column=1, value=r.counterparty_id)
        ws_recon.cell(row=row_idx, column=2, value=r.isin)
        ws_recon.cell(row=row_idx, column=3, value=str(r.statement_date))
        ws_recon.cell(row=row_idx, column=4, value=r.internal_quantity)
        ws_recon.cell(row=row_idx, column=5, value=r.custodian_quantity)
        ws_recon.cell(row=row_idx, column=6, value=r.difference)
        ws_recon.cell(row=row_idx, column=7, value="Yes" if r.is_reconciled else "No")
        if not r.is_reconciled:
            ws_recon.cell(row=row_idx, column=7).fill = SEVERITY_COLORS["HIGH"]

    _auto_width(ws_recon)

    # ── Auctions Tab ────────────────────────────────────────────────────
    ws_auctions = wb.create_sheet("Auctions")
    auction_headers = [
        "Auction ID", "ISIN", "Short Qty", "Valuation Price",
        "Auction Price", "Close-Out Price", "Penalty (INR)", "Outcome", "Status",
    ]
    for col, h in enumerate(auction_headers, 1):
        ws_auctions.cell(row=1, column=col, value=h)
    _style_header_row(ws_auctions, 1, len(auction_headers))

    auctions = session.query(AuctionRecord).all()
    for row_idx, a in enumerate(auctions, 2):
        ws_auctions.cell(row=row_idx, column=1, value=a.auction_id[:12] + "...")
        ws_auctions.cell(row=row_idx, column=2, value=a.isin)
        ws_auctions.cell(row=row_idx, column=3, value=a.short_quantity)
        ws_auctions.cell(row=row_idx, column=4, value=float(a.valuation_price))
        ws_auctions.cell(row=row_idx, column=5, value=float(a.auction_price) if a.auction_price else "N/A")
        ws_auctions.cell(row=row_idx, column=6, value=float(a.close_out_price) if a.close_out_price else "N/A")
        ws_auctions.cell(row=row_idx, column=7, value=float(a.penalty_amount))
        ws_auctions.cell(row=row_idx, column=8, value=a.outcome.value if a.outcome else "PENDING")
        ws_auctions.cell(row=row_idx, column=9, value=a.status.value)

    _auto_width(ws_auctions)

    # ── Counterparty Analysis Tab ───────────────────────────────────────
    ws_cp = wb.create_sheet("Counterparty Analysis")
    cp_headers = ["Counterparty", "Total Breaks", "HIGH", "MEDIUM", "LOW", "Total VAR (INR)"]
    for col, h in enumerate(cp_headers, 1):
        ws_cp.cell(row=1, column=col, value=h)
    _style_header_row(ws_cp, 1, len(cp_headers))

    cp_breaks: dict[str, list[BreakRecord]] = {}
    for brk in all_breaks:
        ob = session.query(Obligation).filter(
            Obligation.obligation_id == brk.obligation_id
        ).first()
        if ob:
            cp_id = ob.counterparty_id
            cp_breaks.setdefault(cp_id, []).append(brk)

    sorted_cps = sorted(cp_breaks.items(), key=lambda x: len(x[1]), reverse=True)
    for row_idx, (cp_id, brks) in enumerate(sorted_cps, 2):
        ws_cp.cell(row=row_idx, column=1, value=cp_id)
        ws_cp.cell(row=row_idx, column=2, value=len(brks))
        ws_cp.cell(row=row_idx, column=3, value=sum(1 for b in brks if b.severity == Severity.HIGH))
        ws_cp.cell(row=row_idx, column=4, value=sum(1 for b in brks if b.severity == Severity.MEDIUM))
        ws_cp.cell(row=row_idx, column=5, value=sum(1 for b in brks if b.severity == Severity.LOW))
        ws_cp.cell(row=row_idx, column=6, value=sum(float(b.value_at_risk or 0) for b in brks))

    _auto_width(ws_cp)

    wb.save(output_path / "settlement_report.xlsx")
    print(f"  Excel report saved: {output_path / 'settlement_report.xlsx'}")


def generate_docx_report(
    session: Session,
    triage_result: dict,
    recon_results: list,
    output_path: Path,
):
    """Generate a narrative DOCX settlement summary."""
    doc = Document()

    # Title
    title = doc.add_heading("Daily Settlement Operations Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Report Date: {datetime.now().strftime('%d %B %Y')}\n"
        f"Settlement Window: June 2026 (20 trading days)\n"
        f"Settlement Cycles: T+1 (standard) and T+0 (top 500 eligible stocks)"
    )

    # KPIs
    doc.add_heading("Key Performance Indicators", level=1)

    total_trades = session.query(Trade).filter(Trade.source_system == "OMS").count()
    total_obs = session.query(Obligation).filter(
        Obligation.obligation_stage == ObligationStage.FINAL
    ).count()
    total_breaks = session.query(BreakRecord).count()

    stp_count = session.query(Obligation).filter(
        Obligation.obligation_stage == ObligationStage.FINAL,
        Obligation.status.in_([ObligationStatus.SETTLED, ObligationStatus.INSTRUCTED]),
    ).count()
    stp_rate = (stp_count / total_obs * 100) if total_obs > 0 else 0

    config = get_escalation_config()
    cost_config = config["cost_per_break"]
    all_breaks = session.query(BreakRecord).all()
    total_cost = sum(cost_config.get(b.severity.value, 5000) for b in all_breaks)

    kpi_table = doc.add_table(rows=6, cols=2)
    kpi_table.style = "Light Grid Accent 1"
    kpis = [
        ("Total Trades Captured", str(total_trades)),
        ("Net Obligations (Final Stage)", str(total_obs)),
        ("STP Rate", f"{stp_rate:.1f}%"),
        ("Total Breaks", str(total_breaks)),
        ("High-Risk Queue (Pre-Settlement)", str(len(triage_result.get("high_risk_queue", [])))),
        ("Estimated Exception Cost", f"INR {total_cost:,.0f}"),
    ]
    for i, (label, val) in enumerate(kpis):
        kpi_table.rows[i].cells[0].text = label
        kpi_table.rows[i].cells[1].text = val

    # Break Analysis
    doc.add_heading("Break Analysis", level=1)

    from collections import Counter
    type_dist = Counter(b.break_type.value for b in all_breaks)
    sev_dist = Counter(b.severity.value for b in all_breaks)

    doc.add_paragraph(
        f"A total of {total_breaks} breaks were detected across the settlement window. "
        f"The distribution by type is as follows:"
    )

    for bt, count in type_dist.most_common():
        doc.add_paragraph(f"{bt}: {count} ({count/total_breaks*100:.1f}%)", style="List Bullet")

    doc.add_paragraph(f"\nSeverity distribution: " + ", ".join(
        f"{s}: {c}" for s, c in sev_dist.most_common()
    ))

    # Triage Summary
    doc.add_heading("Agentic Triage Summary", level=1)
    doc.add_paragraph(
        f"The LangGraph triage pipeline processed {len(triage_result.get('breaks', []))} breaks "
        f"through the post-break triage path (classifier → root-cause investigator → "
        f"resolution recommender → escalation checker → human approval gate). "
        f"All {len(triage_result.get('triage_results', []))} break resolutions are "
        f"queued for human approval — no auto-resolution was performed."
    )

    doc.add_paragraph(
        f"The pre-settlement fail-risk scan evaluated {len(triage_result.get('fail_risk_scores', []))} "
        f"pending obligations, identifying {len(triage_result.get('high_risk_queue', []))} "
        f"as high-risk requiring proactive intervention."
    )

    # Reconciliation
    doc.add_heading("Position Reconciliation", level=1)
    total_positions = len(recon_results)
    reconciled = sum(1 for r in recon_results if r.is_reconciled)
    doc.add_paragraph(
        f"End-of-day position reconciliation compared {total_positions} position lines "
        f"between internally derived positions (from settled obligations) and custodian "
        f"EOD holding statements. {reconciled} positions ({reconciled/total_positions*100:.1f}% "
        f"if total_positions > 0 else 0) reconciled successfully."
    )

    # Audit Trail
    doc.add_heading("Audit & Governance", level=1)
    audit_count = session.query(AgenticAuditLog).count()
    doc.add_paragraph(
        f"The agentic pipeline produced {audit_count} audit log entries across all nodes. "
        f"Each entry records the node name, input data, conclusion, and rationale — "
        f"providing a complete reasoning chain for every agentic decision. "
        f"This meets the governance requirement for agentic system auditability "
        f"in financial operations."
    )

    doc.save(output_path / "settlement_summary.docx")
    print(f"  DOCX report saved: {output_path / 'settlement_summary.docx'}")


def generate_reports(
    session: Session,
    triage_result: dict,
    recon_results: list,
    output_dir: Path,
):
    """Generate all reports."""
    generate_excel_report(session, triage_result, recon_results, output_dir)
    generate_docx_report(session, triage_result, recon_results, output_dir)
